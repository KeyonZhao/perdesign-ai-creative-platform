#!/usr/bin/env python3
import argparse
import json
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def extract_docx(path: Path):
    try:
        from docx import Document
        doc = Document(path)
        blocks = []
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:
                blocks.append({"kind": "paragraph", "index": i, "style": p.style.name, "text": text})
        tables = []
        for ti, table in enumerate(doc.tables):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append({"kind": "table", "index": ti, "rows": rows})
        return blocks, tables
    except ImportError:
        with zipfile.ZipFile(path) as zf:
            root = ET.fromstring(zf.read("word/document.xml"))
        blocks = []
        for i, p in enumerate(root.iter(W + "p")):
            text = "".join(t.text or "" for t in p.iter(W + "t")).strip()
            if text:
                blocks.append({"kind": "paragraph", "index": i, "style": "Unknown", "text": text})
        return blocks, []


def extract_pdf(path: Path):
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("PDF extraction requires pdfplumber from the workspace runtime") from exc
    blocks, tables = [], []
    with pdfplumber.open(path) as pdf:
        for pi, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            for li, line in enumerate(text.splitlines()):
                line = line.strip()
                if line:
                    blocks.append({"kind": "paragraph", "page": pi, "index": li, "style": "PDF", "text": line})
            for ti, table in enumerate(page.extract_tables() or []):
                tables.append({"kind": "table", "page": pi, "index": ti, "rows": [[c or "" for c in row] for row in table]})
    return blocks, tables


def to_markdown(source, blocks, tables):
    lines = [f"# Extracted source: {source.name}", "", f"Paragraphs: {len(blocks)}", f"Tables: {len(tables)}", ""]
    for block in blocks:
        text = block["text"]
        style = block.get("style", "")
        if style.startswith("Heading"):
            match = re.search(r"(\d+)$", style)
            level = min(int(match.group(1)) if match else 2, 6)
            lines.extend(["#" * level + " " + text, ""])
        else:
            location = f"P{block.get('index')}" if "page" not in block else f"Page {block['page']} · L{block.get('index')}"
            lines.extend([f"<!-- {location} · {style} -->", text, ""])
    for table in tables:
        loc = f"Table {table['index']}" if "page" not in table else f"Page {table['page']} · Table {table['index']}"
        lines.extend([f"## {loc}", ""])
        for row in table["rows"]:
            lines.append(" | ".join(str(c).replace("\n", " / ") for c in row))
        lines.append("")
    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(description="Extract complete DOCX/PDF content for strategic map analysis")
    parser.add_argument("input")
    parser.add_argument("--out", required=True)
    parser.add_argument("--markdown")
    args = parser.parse_args()
    source = Path(args.input).expanduser().resolve()
    if not source.exists():
        raise SystemExit(f"Input not found: {source}")
    suffix = source.suffix.lower()
    if suffix == ".docx":
        blocks, tables = extract_docx(source)
    elif suffix == ".pdf":
        blocks, tables = extract_pdf(source)
    else:
        raise SystemExit("Supported formats: .docx, .pdf")
    payload = {"source": str(source), "paragraph_count": len(blocks), "table_count": len(tables), "blocks": blocks, "tables": tables}
    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    md_path = Path(args.markdown) if args.markdown else out.with_suffix(".md")
    md_path.write_text(to_markdown(source, blocks, tables), encoding="utf-8")
    print(json.dumps({"json": str(out), "markdown": str(md_path), "paragraphs": len(blocks), "tables": len(tables)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
