from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).parent / "策划案到可编辑战略思维导图_Codex执行规范.docx"
NAVY = "173F3C"
TEAL = "1F756A"
MINT = "DDF1EC"
PALE = "F2F7F5"
INK = "18302F"
MUTED = "657977"
WHITE = "FFFFFF"
GOLD = "B88B39"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_bottom_border(paragraph, color=NAVY, size="14"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_font(run, size=None, bold=None, color=INK, name="Arial Unicode MS"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", size=10.5, bold=False, color=INK, before=0, after=6, line=1.2, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, size=10.3, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, size=10.3, color=INK)
    else:
        r = p.add_run(text)
        set_font(r, size=10.3, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.18
    set_font(p.add_run(text), size=10.3, color=INK)
    return p


def add_callout(doc, label, text, fill=MINT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.25)
    cell = table.cell(0, 0)
    cell.width = Inches(6.25)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    set_font(r, size=9, bold=True, color=TEAL)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    set_font(p2.add_run(text), size=11, bold=True, color=NAVY)
    add_para(doc, "", size=2, after=4)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Inches(w)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(h), size=9.3, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, (value, w) in enumerate(zip(row, widths)):
            cells[i].width = Inches(w)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri % 2:
                set_cell_shading(cells[i], PALE)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            set_font(p.add_run(value), size=9.1, color=INK)
    add_para(doc, "", size=2, after=4)
    return table


def heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(p)
    return p


def page_break(doc):
    doc.add_page_break()


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.82)
sec.bottom_margin = Inches(0.78)
sec.left_margin = Inches(0.95)
sec.right_margin = Inches(0.95)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial Unicode MS"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2
for name in ("List Bullet", "List Bullet 2", "List Number"):
    s = styles[name]
    s.font.name = "Arial Unicode MS"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    s.font.size = Pt(10.3)
for level, size, color, before, after in ((1, 17, NAVY, 16, 8), (2, 13, TEAL, 12, 6), (3, 11, NAVY, 8, 4)):
    s = styles[f"Heading {level}"]
    s.font.name = "Arial Unicode MS"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

# Running header/footer
hp = sec.header.paragraphs[0]
hp.text = "CODEX EXECUTION STANDARD  ·  STRATEGIC MIND MAP"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(hp.runs[0], size=8, bold=True, color=MUTED)
add_bottom_border(hp, color="B8CBC7", size="4")
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(fp.add_run("策划案 → 论证体系 → 可编辑导图网页"), size=8, color=MUTED)

# Cover
add_para(doc, "CODEX 执行规范", size=10, bold=True, color=GOLD, before=36, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "从策划案文档到\n可编辑战略思维导图网页", size=27, bold=True, color=NAVY, after=10, line=1.12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "一份可直接上传给空白 Codex 的工作标准与验收协议", size=13, color=TEAL, after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
add_bottom_border(doc.paragraphs[-1], color=TEAL, size="8")
add_para(doc, "", size=5, after=22)
add_callout(doc, "本文件的用途", "将本文件与任意策划案 DOCX / PDF 一起上传给 Codex。Codex 必须先完整理解策划案，再把章节素材重构为一条可论证、可归纳、可交互的战略思维导图，而不是复制目录。")
add_para(doc, "适用输入", size=9, bold=True, color=MUTED, before=18, after=4)
add_para(doc, "品牌策划案、产品创新方案、研究报告、概念提案、市场洞察、用户研究、商业计划及其他长篇结构化文档。", size=11, after=12)
add_para(doc, "目标输出", size=9, bold=True, color=MUTED, after=4)
add_para(doc, "一个可在线或本地预览、可拖拽平移、可缩放、可展开收起、可搜索、可编辑节点、可增删层级并可自动保存的思维导图网页。", size=11, after=12)
add_para(doc, "版本 1.0  ·  复现标准：海信空气感知屏战略思维导图", size=8.5, color=MUTED, before=22, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break(doc)
heading(doc, "0. 给 Codex 的直接执行指令", 1)
add_callout(doc, "最高优先级指令", "你正在执行“策划案 → 可编辑战略思维导图网页”任务。本规范是交付标准。除非用户明确要求改变，否则必须完整遵守。不要只总结章节，不要先入为主地从产品结论开始，也不要在未读完文档时制作导图。")
heading(doc, "0.1 启动方式", 2)
add_para(doc, "用户通常会同时上传本规范与一份策划案。收到后直接开始工作；仅当缺少策划案文件、文件损坏或输出位置完全不可判断时，才询问用户。")
add_para(doc, "建议用户在空白 Codex 中发送：", bold=True, color=NAVY)
add_callout(doc, "可复制任务语句", "请阅读全文策划案，并严格按照《策划案到可编辑战略思维导图 Codex 执行规范》制作可预览、可编辑的思维导图网页。不要复制目录，要重构完整论证过程，并完成验证。", fill="EEF2F7")
heading(doc, "0.2 最终交付必须同时满足", 2)
for item in [
    "内容层：导图能够解释“为什么得出这个产品定义”，而不只是展示“定义是什么”。",
    "逻辑层：任意子节点都能被父节点准确概括；兄弟节点处于同一分类维度且尽量互斥、共同完整。",
    "结构层：同时呈现前置论证过程和最终创新策略，阶段输出能自然导向下一阶段。",
    "产品层：网页支持平移、缩放、折叠、搜索、节点编辑、增删、自动保存、导入导出与视图复位。",
    "验证层：内容覆盖、父子逻辑、网页构建和核心交互均完成检查后才可交付。",
]: add_bullet(doc, item)

heading(doc, "1. 核心原则：这不是目录可视化", 1)
heading(doc, "1.1 错误输出的典型特征", 2)
add_table(doc, ["错误做法", "为什么不合格", "正确替代"], [
    ("把原文一级、二级标题直接搬成节点", "章节是写作顺序，不一定是推理关系", "按需求、证据、洞察、机会、策略重新归类"),
    ("开篇直接呈现产品定义和功能", "用户看不到结论从何而来", "先展示需求整理、问题研究、人群、场景、竞品和机会收敛"),
    ("节点只是关键词，没有解释", "无法判断归纳关系，也失去原文细节", "标题表达结论，detail 保存证据、边界或例子"),
    ("所有内容都平铺在同一层", "不同抽象层级混杂，无法阅读", "建立阶段 → 主题 → 判断 → 证据/行动的多层结构"),
    ("为了层级多而机械拆词", "节点虽多但没有信息增量", "只有存在独立判断、证据、对象或行动时才拆分"),
], [1.45, 2.35, 2.45])
heading(doc, "1.2 合格导图的本质", 2)
add_para(doc, "合格导图是一套“可导航的论证模型”。它应该允许读者从根节点一路回答：客户为什么提出需求？真实问题是什么？谁受影响？在哪些场景发生？市场如何解决？证据收敛出什么机会？机会如何转化为产品定义、功能、形态、交互与商业闭环？")
add_callout(doc, "一句话判断标准", "如果删除原文目录，读者仍能只看导图理解策划案的主要观点、推导逻辑与执行结论，这张导图才算成立。")

page_break(doc)
heading(doc, "2. 完整工作流程", 1)
add_para(doc, "必须按以下顺序完成。可以内部迭代，但不得跳过阅读全文、逻辑重构和验收。")
for step in [
    "读取文件：提取正文、标题层级、表格、批注、脚注、图片说明及引用；确认页数、段落数和章节范围。",
    "阅读全文：先形成文档级理解，再记录关键判断；不得读到一个结论就提前开工。",
    "建立内容账本：标注每个重要信息单元属于需求、问题、用户、场景、竞品、机会、产品、执行、商业或风险。",
    "重构论证链：把章节素材重排为“研究起点 → 分析证据 → 阶段输出 → 机会收敛 → 产品定义”。",
    "设计树结构：确定根节点、阶段节点、主题节点、判断节点和证据/行动节点；检查父子归纳关系。",
    "编写节点文案：标题短而有判断，说明保留原文事实、例子、边界、风险和应对策略。",
    "制作交互网页：实现完整导图、编辑器、视图控制和本地持久化。",
    "构建与预览：确保页面能够编译并打开；修复实际错误。",
    "执行双重验收：先验收内容逻辑，再验收交互与布局；未通过不得宣称完成。",
]: add_number(doc, step)

heading(doc, "3. 第一阶段：如何完整理解策划案", 1)
heading(doc, "3.1 全文读取要求", 2)
for item in [
    "读取所有非空段落，不能只读取目录、摘要、前几页或标题。",
    "表格必须逐行读取；表格常承载竞品矩阵、配置分级、研究结论和执行计划。",
    "图片无法直接提取文字时，至少记录图片标题、图注、所在章节和它支持的论点。",
    "区分作者的结论、论据、案例、建议、风险、传播话术和重复强调；不要把它们混成同一类节点。",
    "重复内容不重复堆节点，而是合并为一个更强的判断，并在说明中保留多个支持证据。",
]: add_bullet(doc, item)
heading(doc, "3.2 内容账本字段", 2)
add_table(doc, ["字段", "需要记录什么", "用途"], [
    ("原文位置", "章节、标题或段落范围", "覆盖检查与回溯"),
    ("信息类型", "需求 / 问题 / 洞察 / 证据 / 机会 / 策略 / 风险", "重构逻辑阶段"),
    ("核心判断", "用一句完整陈述表达", "转为节点标题"),
    ("支持信息", "数据、案例、现象、用户语言、竞品信号", "转为 detail 或子节点"),
    ("对象", "用户、空间、产品、品牌、渠道、技术", "避免分类维度混乱"),
    ("因果方向", "它解释什么、导向什么", "形成阶段输出"),
    ("优先级", "核心 / 重要 / 补充", "控制默认展开与信息密度"),
], [1.0, 2.25, 3.0])
heading(doc, "3.3 只允许基于文档推导", 2)
add_para(doc, "节点内容优先来自用户上传的策划案。可以进行归纳、合并、改写和逻辑推断，但不能凭空添加市场数据、用户结论、竞品事实或技术指标。若需要外部补充研究，必须先明确其来源，并把“文档观点”和“外部验证”区分开。")

page_break(doc)
heading(doc, "4. 第二阶段：把章节重构为论证过程", 1)
add_para(doc, "默认使用以下七阶段骨架。文档不具备某一阶段时可以合并，但不得因此直接跳到结论。")
add_table(doc, ["阶段", "核心问题", "阶段输出"], [
    ("01 需求整理", "客户说要什么？背后真正想解决什么？有哪些约束？", "被重新定义的研究任务"),
    ("02 问题研究", "品类、价值与现有体验的核心矛盾是什么？", "需要被解决的关键问题"),
    ("03 人群分析", "谁会使用、购买或受影响？他们要完成什么任务？", "用户价值与购买理由"),
    ("04 场景研究", "问题何时、何地、由什么触发？现有方案为何不足？", "能够验证价值的高频情境"),
    ("05 竞品分析", "直接竞品怎么做？跨行业抬高了什么标准？", "市场空白、可借鉴与不可照搬"),
    ("06 机会收敛", "哪些方向同时满足用户、品类、品牌与商业标准？", "被筛选后的产品机会"),
    ("07 产品定义", "产品是什么、为谁、解决什么、如何实现与赚钱？", "功能、形态、交互、路线、风险与执行"),
], [1.25, 3.25, 1.75])
heading(doc, "4.1 每个阶段内部必须形成闭环", 2)
add_para(doc, "一个阶段不应只是若干主题的集合。推荐使用以下内部结构：")
add_bullet(doc, "输入：来自客户、文档事实、用户现象或竞品信号。")
add_bullet(doc, "分析维度：把输入分成互不混杂的观察角度。")
add_bullet(doc, "关键判断：用明确陈述总结每一维度的意义。")
add_bullet(doc, "证据或例子：说明为什么可以得出判断。")
add_bullet(doc, "阶段输出：用一句结论说明本阶段为下一阶段提供了什么。")
heading(doc, "4.2 因果链写法", 2)
add_callout(doc, "推荐句式", "现象 / 需求 → 暴露的问题 → 用户与场景证据 → 竞品参照 → 机会筛选 → 产品定义 → 功能与体验 → 商业价值与风险闭环")

heading(doc, "5. 第三阶段：设计可归纳的树结构", 1)
heading(doc, "5.1 节点层级职责", 2)
add_table(doc, ["层级", "职责", "文案示例"], [
    ("L0 根节点", "定义整个研究对象与最终任务", "某品牌产品创新论证全景"),
    ("L1 阶段节点", "表达研究过程或战略模块", "03｜人群分析：谁会为它买单"),
    ("L2 主题节点", "一个阶段下的分析维度", "家庭关系带来的关注"),
    ("L3 判断节点", "可独立成立的观点", "有孩子家庭关注洁净、温差与直吹"),
    ("L4+ 证据/行动", "例子、指标、触发、策略、功能或对策", "显示活动区、舒适区与人体避风范围"),
], [1.05, 2.35, 2.85])
heading(doc, "5.2 父子关系四问", 2)
for item in [
    "归纳问：父节点能否完整概括所有直接子节点？",
    "同维问：兄弟节点是否使用同一分类标准，例如都是人群、场景、问题或策略？",
    "增量问：子节点是否增加了独立信息，而非重复父节点换一种说法？",
    "必要问：如果删除这个节点，论证或执行信息是否会明显损失？",
]: add_bullet(doc, item)
heading(doc, "5.3 拆分与合并规则", 2)
add_bullet(doc, "当一句话同时包含两个不同对象、两种因果或两个执行动作时，应拆成兄弟节点。")
add_bullet(doc, "当多个原文章节反复支持同一结论时，应合并为一个判断，把各处证据放进说明或下级节点。")
add_bullet(doc, "列表元素不足以独立解释、没有信息增量时，不为追求层级而继续拆分。")
add_bullet(doc, "核心内容应达到 4–6 层；补充信息可以较浅。层级深度服务理解，不追求统一。")

page_break(doc)
heading(doc, "6. 节点文案标准", 1)
heading(doc, "6.1 标题写结论，不只写名词", 2)
add_table(doc, ["弱标题", "强标题"], [
    ("用户需求", "用户购买的是“被照顾的确定感”，不是屏幕"),
    ("行业趋势", "家电屏幕化的本质是角色升级，不是尺寸变大"),
    ("竞品", "直接竞品能力趋同，但价值表达仍弱"),
    ("卧室", "卧室屏幕必须让位于睡眠"),
    ("风险", "信息不准确会迅速损害用户信任"),
], [2.5, 3.75])
heading(doc, "6.2 标题与说明的分工", 2)
add_bullet(doc, "title：短、可扫描、包含判断；推荐 8–28 个汉字。")
add_bullet(doc, "detail：保存原因、证据、例子、边界和执行要求；推荐 1–3 句。")
add_bullet(doc, "tag：只标记节点角色，例如“研究起点”“用户洞察”“推导结果”，不要替代标题。")
heading(doc, "6.3 信息密度", 2)
add_bullet(doc, "一级阶段通常 5–8 个；每个阶段 3–6 个主题；每个主题 2–6 个判断。")
add_bullet(doc, "长文档通常需要 80–180 个有效节点。节点多不等于质量高，必须通过父子关系检查。")
add_bullet(doc, "页面第一眼只展开 L0–L1 或少量关键 L2；其余通过折叠逐步阅读。")

heading(doc, "7. 网页产品标准", 1)
heading(doc, "7.1 必做能力", 2)
add_table(doc, ["能力", "交互要求", "验收方式"], [
    ("拖拽平移", "按住画布空白处拖动；节点和表单操作不触发平移", "鼠标与触摸均可移动视角"),
    ("缩放", "提供放大、缩小和百分比反馈", "大导图可查看全局与细节"),
    ("适应视图", "一键重置缩放和平移", "操作后回到初始可读位置"),
    ("展开/收起", "有子节点时显示明确控件；搜索时自动暴露匹配路径", "各阶段可独立折叠"),
    ("搜索", "同时搜索标题与说明并高亮结果", "输入关键词可定位观点"),
    ("节点编辑", "可编辑标题、说明和标签", "修改后立即反映在导图"),
    ("结构编辑", "可添加子节点、删除非根节点", "树结构可持续扩展"),
    ("本地保存", "使用浏览器本地存储保存用户修改", "刷新页面后内容仍在"),
    ("导入导出", "以 JSON 等结构化格式备份与恢复", "导出的数据可再次导入"),
], [1.15, 3.05, 2.05])
heading(doc, "7.2 视觉与布局", 2)
for item in [
    "必须有清晰连接线，不能让节点像互不相关的卡片堆。",
    "根节点、阶段节点、主题节点和普通节点应通过色彩、尺寸或边框形成层级差异。",
    "采用克制的商务视觉；优先保证长时间阅读，不使用无意义渐变、炫光或大面积装饰。",
    "画布应足够大且不被页面滚动限制；编辑面板应固定在侧边并可关闭。",
    "桌面端为主要体验，同时保证窄屏不溢出、编辑器不会完全遮挡内容。",
    "默认折叠策略应让读者先看到完整论证阶段，再自行深入。",
]: add_bullet(doc, item)
heading(doc, "7.3 推荐数据模型", 2)
add_callout(doc, "节点结构", "MapNode = { id: string; title: string; detail?: string; tag?: string; children?: MapNode[] }", fill="EEF2F7")
add_para(doc, "实现技术可以根据工作区现状选择。优先复用现有框架和依赖；不要为了导图引入沉重依赖。如果现有项目没有图形库，可以使用嵌套 DOM、CSS 连接线与 transform 完成平移缩放。")

page_break(doc)
heading(doc, "8. 内容覆盖与质量验收", 1)
heading(doc, "8.1 内容覆盖检查", 2)
for item in [
    "策划案的核心结论是否全部进入导图？",
    "重要论据、用户矛盾、场景、竞品信号、功能建议、风险与对策是否被保留？",
    "是否有大段内容只因不在标题中而被遗漏？",
    "重复观点是否被合并，而非机械重复？",
    "导图是否既保留作者观点，又通过重组提高了逻辑清晰度？",
]: add_bullet(doc, item)
heading(doc, "8.2 逻辑验收", 2)
for item in [
    "根节点能概括全部一级阶段。",
    "每个阶段都有明确研究问题和阶段输出。",
    "每个父节点都能概括直接子节点，不存在明显错挂。",
    "兄弟节点处于同一分析维度，没有“人群、功能、风险”混在一层。",
    "产品定义能追溯到前面的需求、问题、人群、场景和竞品证据。",
    "功能不是凭空出现，而是对应某个问题、任务或场景。",
    "风险都有对应的控制策略，商业价值有实现路径。",
]: add_bullet(doc, item)
heading(doc, "8.3 网页验收", 2)
for item in [
    "页面成功构建，无阻断性交互错误。",
    "平移、缩放和复位连续可用，不与节点点击冲突。",
    "折叠、搜索、编辑、增删、自动保存、导入和导出均可工作。",
    "默认视图能看清总体阶段，不会一打开就铺满全部细节。",
    "长标题能够换行，节点不重叠，连接关系可辨认。",
]: add_bullet(doc, item)
heading(doc, "8.4 评分门槛", 2)
add_table(doc, ["维度", "权重", "通过标准"], [
    ("全文覆盖", "25%", "主要观点、证据、场景、策略和风险无重大遗漏"),
    ("论证完整", "25%", "结论可从前置分析自然推导，不先给答案"),
    ("父子逻辑", "20%", "绝大多数节点可被父级准确归纳"),
    ("内容表达", "10%", "标题有判断，说明有证据，避免空泛名词"),
    ("交互完整", "15%", "全部必做交互可用且不冲突"),
    ("视觉可读", "5%", "层级清楚、默认视图合理、长内容可阅读"),
], [1.55, 0.8, 3.85])
add_callout(doc, "交付门槛", "总分必须达到 85/100，且“论证完整”“父子逻辑”“交互完整”三项不得低于各自权重的 80%。")

heading(doc, "9. 常见失败与修正动作", 1)
add_table(doc, ["失败现象", "修正动作"], [
    ("打开就是产品定义", "在原有内容前增加需求、问题、人群、场景、竞品和机会收敛；保留原内容作为最终阶段"),
    ("导图像目录", "隐藏原章节编号，按信息类型与因果关系重新归类"),
    ("层级很深但内容空", "合并只有措辞差异的节点；每个叶子节点必须增加证据或行动信息"),
    ("节点太长", "标题保留判断，背景和例子移入 detail"),
    ("展开后无法浏览", "默认折叠一级阶段，增加平移、缩放和适应视图"),
    ("拖拽会误点节点", "只允许从空白画布启动平移，交互控件阻止拖拽"),
    ("刷新后修改消失", "使用版本化本地存储键；结构更新时处理迁移或升级键"),
    ("为了发布破坏现有项目", "保留现有架构和首页，使用独立路由；发布失败不应损害本地交付"),
], [2.25, 4.0])

page_break(doc)
heading(doc, "10. 可直接执行的完整任务协议", 1)
add_para(doc, "以下内容是本规范的压缩执行版。Codex 读完本文件后，应把它视为当前任务的行动协议。", color=MUTED)
protocol = [
    "完整读取用户上传的策划案正文、标题、表格与图注，先形成文档级理解。",
    "不要复制原文目录。建立内容账本，把信息归类为需求、问题、人群、场景、竞品、机会、产品、执行、商业和风险。",
    "用七阶段结构重构内容：需求整理 → 问题研究 → 人群分析 → 场景研究 → 竞品分析 → 机会收敛 → 产品定义与创新策略。",
    "每个阶段必须包含分析维度、关键判断、支持信息和阶段输出；产品结论必须能追溯到前置证据。",
    "设计多层树结构。所有子节点必须可被父节点归纳；兄弟节点必须处于同一分类维度。核心内容应达到 4–6 层，但不得机械拆分。",
    "节点标题写判断，说明写原因、证据、例子、边界或行动。充分覆盖原文，不凭空添加事实。",
    "制作独立、可预览的网页导图，并实现：空白画布拖拽平移、触摸平移、缩放、适应视图、展开收起、搜索高亮、节点编辑、添加子节点、删除节点、本地自动保存、JSON 导入导出和恢复原版。",
    "默认只展开总体论证阶段，让读者先理解全局，再逐层深入。保留清晰连接线和层级化视觉。",
    "完成构建检查，并按内容覆盖、论证完整、父子逻辑、交互完整和视觉可读五方面验收。",
    "最终向用户交付可打开的预览地址或本地页面，并简要说明论证阶段与编辑能力。不得在核心功能未验证时宣称完成。",
]
for item in protocol: add_number(doc, item)

heading(doc, "11. 交付前最终清单", 1)
checks = [
    "已完整阅读全文与表格，而不是只看标题。",
    "导图开头是研究与论证过程，不是创新结论。",
    "需求、问题、人群、场景、竞品、机会和产品定义均有独立阶段。",
    "每个阶段都有清晰的阶段输出。",
    "产品定义能从机会收敛自然得出。",
    "现有策划案中的关键观点、例子、功能、风险和对策已被覆盖。",
    "不存在明显错挂节点或混合分类维度。",
    "默认视图可读，各阶段可以展开收起。",
    "画布可以用鼠标和触摸拖动平移。",
    "缩放与适应视图可用。",
    "搜索、编辑、增删、保存、导入导出可用。",
    "构建成功，页面能够打开。",
]
for item in checks:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run("☐  " + item), size=10.3, color=INK)

add_callout(doc, "完成定义", "只有当以上清单全部满足，才可以向用户汇报“已完成”。如果存在阻塞，应明确说明缺失项、已完成部分和下一步，不得用“页面已生成”替代质量验收。")

# Appendix
page_break(doc)
heading(doc, "附录 A｜建议的导图顶层模板", 1)
for title, children in [
    ("01｜需求整理", "显性需求 / 隐性诉求 / 约束条件 / 阶段输出"),
    ("02｜问题研究", "品类矛盾 / 价值矛盾 / 界面矛盾 / 失败假设 / 阶段输出"),
    ("03｜人群分析", "基础画像 / 家庭角色 / 用户任务 / 使用张力 / 阶段输出"),
    ("04｜场景研究", "场景触发 / 用户问题 / 现有不足 / 验证价值 / 阶段输出"),
    ("05｜竞品分析", "直接竞品 / 跨行业参照 / 可借鉴 / 不可照搬 / 品牌机会 / 阶段输出"),
    ("06｜机会收敛", "筛选标准 / 核心推导链 / 排除方向 / 产品机会 / 阶段输出"),
    ("07｜产品定义与策略", "产品角色 / 核心价值 / 功能体系 / 交互 / 形态 / 产品组合 / 商业 / 风险 / 执行"),
]:
    p = add_para(doc, title, size=11, bold=True, color=NAVY, before=7, after=2)
    add_para(doc, children, size=9.5, color=MUTED, after=4)

heading(doc, "附录 B｜Codex 最终回复模板", 1)
add_callout(doc, "回复格式", "已完成策划案的全文理解与逻辑重构，并制作可编辑思维导图网页。导图按照“需求整理—问题研究—人群分析—场景研究—竞品分析—机会收敛—产品定义与策略”展开，共包含 [节点数量] 个观点节点。支持拖拽平移、缩放、折叠、搜索、节点编辑、增删、自动保存与导入导出。预览地址：[地址]。", fill="EEF2F7")

doc.core_properties.title = "策划案到可编辑战略思维导图：Codex执行规范"
doc.core_properties.subject = "将长篇策划案转化为具备完整论证过程的可编辑思维导图网页"
doc.core_properties.author = "Codex"
doc.core_properties.keywords = "Codex, 策划案, 思维导图, 战略分析, 网页, 执行规范"
doc.save(OUT)
print(OUT)
